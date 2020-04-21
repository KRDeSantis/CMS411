"""
This was made to test two features of docx; One being the heading function, which is commented at the bottom
The other was finding the bolding of specific parts of the whole system

Neither were found to work well in trying to prune down sections of documents to more info heavy parts, 
so aside from bolding where the problems were found after trying to incorporate it into our then solution
was abandoned.
"""


from docx import Document
from docx.shared import Pt
import re
#Copy of Sources Sought Synopsis Manuals 8 Jan 2020
#02 - DRAFT_6_SUAS_Engineering_Services_Sample_TO_08Nov19 (Task Order 0
document2 = Document("Copy of Sources Sought Synopsis Manuals 8 Jan 2020.docx")


runtext = []
boldingFound = False
breakOut = False
#For each run in the paragraph, look for the bolded word "Requirements"
#Once found, print out everything up until the next bolded heading.
for para in document2.paragraphs:
    if(breakOut):
        break
    for runs in para.runs:
        if(boldingFound == False):
            #Use regular expressions to find the text we want in the run
            requirementSearch = re.findall("Requirements", runs.text)
            #if what we get back has found the text we're looking for and the run is bolded
            #Then we add this sentence's text and the text until the next bolded heading
            if(len(requirementSearch) >= 1 and runs.bold):
                runtext.append(runs.text)
                boldingFound = True
        #if we find a new bolded run, in theory we've found everything the section we're looking through has
        #to offer. Otherwise, keep adding text.
        else:
            if(runs.bold):
                breakOut = True
                break
            else:
                runtext.append(runs.text)
         
print(runtext)


'''
underHeading = False
for par in document2.paragraphs:
    if (underHeading == True):
        print(par.text + "\n")
        underHeading = False
    if (par.style.name.startswith('Heading') ):
        underHeading = True


make it so that until next heading grab all the info
search for headings with our keywords
'''

# code is pretty simple but there are many issues that are headaches:
# some documents have alot of headings that are random/ not needed
# making a list of headings that are acceptable and checking them might help, but it also restricts how useful the program is
# (any doc without any of those headings would not have a summary )
# in some documents there is and empty section between the heading and the next text block, while in others it is right after
#Additionally, not all documents even use headings for the document whatsoever